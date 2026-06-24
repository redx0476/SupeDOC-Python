"""Convert proposal HTML (section_html / description) into Word content.

Builds the HTML into a scratch ``python-docx`` Document; the merger then moves
those paragraphs/tables into the rendered document in place of a marker. Lists
are flattened to bullet/number paragraphs (portable across templates) and HTML
tables become real Word tables, mirroring the Docmosis output.
"""
from __future__ import annotations

import html as html_lib
import re

from docx import Document
from docx.shared import Pt
from lxml import html as lxml_html

_BULLET = "\u2022\u00a0"  # "• " (bullet + non-breaking space)
_INLINE_BOLD = {"b", "strong"}
_INLINE_ITALIC = {"i", "em"}


def build_html_doc(raw_html):
    """Return a scratch ``Document`` containing ``raw_html`` as Word content."""
    doc = Document()
    text = (raw_html or "").strip()
    if text:
        for node in _top_level_nodes(text):
            _render_block(doc, node)
    return doc


def _top_level_nodes(text):
    root = lxml_html.fragment_fromstring(text, create_parent="div")
    children = list(root)
    if root.text and root.text.strip():
        p = lxml_html.fromstring(f"<p>{html_lib.escape(root.text)}</p>")
        children = [p] + children
    return children


def _render_block(doc, node):
    tag = (node.tag or "").lower() if isinstance(node.tag, str) else ""

    if tag in ("p", "div"):
        para = doc.add_paragraph()
        _render_inline(para, node)
        _tail_paragraph(doc, node)
    elif tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        para = doc.add_paragraph()
        run = para.add_run(_text_of(node))
        run.bold = True
        run.font.size = Pt(14 - (int(tag[1]) - 1))
        _tail_paragraph(doc, node)
    elif tag == "ul":
        _render_list(doc, node, ordered=False)
        _tail_paragraph(doc, node)
    elif tag == "ol":
        _render_list(doc, node, ordered=True)
        _tail_paragraph(doc, node)
    elif tag == "table":
        _render_table(doc, node)
        _tail_paragraph(doc, node)
    elif tag in ("b", "strong", "i", "em", "span", "a"):
        para = doc.add_paragraph()
        _render_inline(para, node, parent_tag=tag)
    else:
        txt = _text_of(node)
        if txt:
            doc.add_paragraph(txt)


def _tail_paragraph(doc, node):
    if node.tail and node.tail.strip():
        doc.add_paragraph(_collapse_ws(html_lib.unescape(node.tail)))


def _render_list(doc, node, ordered):
    idx = 0
    for li in node.findall("li"):
        idx += 1
        prefix = f"{idx}. " if ordered else _BULLET
        para = doc.add_paragraph()
        para.add_run(prefix)
        _render_inline(para, li)


def _render_inline(para, node, parent_tag=None):
    bold = parent_tag in _INLINE_BOLD if parent_tag else False
    italic = parent_tag in _INLINE_ITALIC if parent_tag else False

    if node.text:
        _add_run(para, node.text, bold, italic)

    for child in node:
        ctag = (child.tag or "").lower() if isinstance(child.tag, str) else ""
        if ctag == "br":
            para.add_run().add_break()
        else:
            cbold = bold or ctag in _INLINE_BOLD
            citalic = italic or ctag in _INLINE_ITALIC
            if child.text:
                _add_run(para, child.text, cbold, citalic)
            for gchild in child:  # one nesting level, e.g. <b><i>..</i></b>
                gtag = (gchild.tag or "").lower() if isinstance(gchild.tag, str) else ""
                if gchild.text:
                    _add_run(para, gchild.text, cbold or gtag in _INLINE_BOLD,
                             citalic or gtag in _INLINE_ITALIC)
                if gchild.tail:
                    _add_run(para, gchild.tail, cbold, citalic)
        if child.tail:
            _add_run(para, child.tail, bold, italic)


def _add_run(para, text, bold, italic):
    text = _preserve_edges(html_lib.unescape(text))
    if not text:
        return
    run = para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _render_table(doc, node):
    rows = node.findall(".//tr")
    if not rows:
        return
    ncols = 0
    for r in rows:
        cells = [c for c in r if isinstance(c.tag, str) and c.tag.lower() in ("th", "td")]
        ncols = max(ncols, len(cells))
    if ncols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=ncols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for r_i, tr in enumerate(rows):
        cells = [c for c in tr if isinstance(c.tag, str) and c.tag.lower() in ("th", "td")]
        for c_i, cell in enumerate(cells):
            if c_i >= ncols:
                break
            wc = table.cell(r_i, c_i)
            wc.text = ""
            run = wc.paragraphs[0].add_run(_text_of(cell))
            if (cell.tag or "").lower() == "th":
                run.bold = True


def _text_of(node):
    return _collapse_ws(html_lib.unescape(node.text_content()))


def _collapse_ws(text):
    return re.sub(r"\s+", " ", text or "").strip()


def _preserve_edges(text):
    """Collapse internal whitespace but keep a single leading/trailing space so
    inter-run spacing (e.g. ``Hello <b>bold</b> world``) is not lost."""
    if not text:
        return ""
    lead = " " if text[:1].isspace() else ""
    trail = " " if text[-1:].isspace() else ""
    core = re.sub(r"\s+", " ", text).strip()
    if not core:
        return " "
    return f"{lead}{core}{trail}"
