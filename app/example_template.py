"""Build the example .docx template users download, restyle in Word, re-upload.

It is a normal Word document whose text contains docxtemplater-style **docxtpl
(Jinja2)** tags. Loop/condition controls use docxtpl's paragraph tag ``{%p ...%}``
so the control paragraphs are removed cleanly at merge time. Rich body content is
injected at ``{{ ... .body }}`` placeholders (filled with a Subdoc by the merger).

Layout mirrors the Docmosis output: cover page -> table of contents ->
auto-numbered hierarchical sections -> Past Performance -> Key Personnel.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def build_example_template() -> bytes:
    """Return the example template as .docx bytes."""
    doc = Document()
    _cover_page(doc)
    _table_of_contents(doc)
    _body_sections(doc)
    _past_performance(doc)
    _resumes(doc)
    _footer(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _p(doc, text="", style=None, align=None, bold=False, size=None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    return para


def _cover_page(doc):
    _p(doc, "{{ company_name }}", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "{{ title }}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    _p(doc, "Solicitation No. {{ solicitation_number }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "Submitted: {{ date }}", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc)
    _p(doc, "Offeror", bold=True)
    _p(doc, "{{ company_name }}")
    _p(doc, "{{ company_city }}, {{ company_state }} {{ company_zip }}")
    _p(doc)
    _p(doc, "Point of Contact", bold=True)
    _p(doc, "{{ company_contact_person }}")
    _p(doc, "{{ company_contact_email }}")
    _p(doc, "{{ company_contact_phone }}")
    doc.add_page_break()


def _table_of_contents(doc):
    _p(doc, "Table of Contents", style="Heading 1")
    # docxtpl paragraph-loop: one TOC line per section/subsection.
    _p(doc, "{%p for s in sections %}")
    _p(doc, "{{ s.number }}  {{ s.title }}")
    _p(doc, "{%p for ss in s.subsections %}")
    _p(doc, "    {{ ss.number }}  {{ ss.title }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p endfor %}")
    doc.add_page_break()


def _body_sections(doc):
    # Three explicit nesting levels (section -> subsection -> sub-subsection),
    # each heading carrying the matching Word Heading style so a restyled
    # template keeps the user's look.
    _p(doc, "{%p for s in sections %}")
    _p(doc, "{{ s.number }}  {{ s.title }}", style="Heading 1")
    _p(doc, "{{ s.body }}")
    _p(doc, "{%p for ss in s.subsections %}")
    _p(doc, "{{ ss.number }}  {{ ss.title }}", style="Heading 2")
    _p(doc, "{{ ss.body }}")
    _p(doc, "{%p for sss in ss.subsections %}")
    _p(doc, "{{ sss.number }}  {{ sss.title }}", style="Heading 3")
    _p(doc, "{{ sss.body }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p endfor %}")


def _past_performance(doc):
    _p(doc, "{%p if has_past_performance %}")
    _p(doc, "Past Performance", style="Heading 1")
    _p(doc, "{%p for pp in past_performance %}")
    _p(doc, "{{ pp.award_title }}", style="Heading 2")
    _p(doc, "Recipient: {{ pp.award_recipient_name }}    Role: {{ pp.role }}")
    _p(doc, "Contract/Task Order: {{ pp.task_order_number }}    Value: {{ pp.total_contract_value }}")
    _p(doc, "Period of Performance: {{ pp.period_of_performance_start }} to {{ pp.period_of_performance_end }}")
    _p(doc, "NAICS: {{ pp.naics }}    PSC: {{ pp.psc }}    Customer: {{ pp.customer_poc }}")
    _p(doc, "{{ pp.body }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p endif %}")


def _resumes(doc):
    _p(doc, "{%p if has_resumes %}")
    _p(doc, "Key Personnel Resumes", style="Heading 1")
    _p(doc, "{%p for r in resumes %}")
    _p(doc, "{{ r.full_name }} \u2014 {{ r.job_title }}", style="Heading 2")
    _p(doc, "Citizenship: {{ r.citizenship }}")
    _p(doc, "Professional Experience", style="Heading 3")
    _p(doc, "{%p for e in r.professional_experience %}")
    _p(doc, "{{ e.position }}, {{ e.company_name }} ({{ e.start_date }} to {{ e.end_date }})", bold=True)
    _p(doc, "{{ e.body }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "Education", style="Heading 3")
    _p(doc, "{%p for ed in r.education %}")
    _p(doc, "{{ ed.degree }} \u2014 {{ ed.institution_name }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "Technical Skills: {{ r.technical_skills }}")
    _p(doc, "Certifications: {{ r.certifications_and_courses }}")
    _p(doc, "Notable Projects: {{ r.notable_projects }}")
    _p(doc, "{%p endfor %}")
    _p(doc, "{%p endif %}")


def _footer(doc):
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.text = ""
    run = para.add_run("{{ company_name }} \u2014 {{ title }}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    import sys

    out = sys.argv[1] if len(sys.argv) > 1 else "example_template.docx"
    with open(out, "wb") as fh:
        fh.write(build_example_template())
    print(f"wrote {out}")
