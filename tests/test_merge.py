"""End-to-end tests for the Python document generator."""
import io

import docx

from app.example_template import build_example_template
from app.merge import merge

SAMPLE = {
    "metadata": {"title": "Test Proposal", "company_name": "Acme Corp"},
    "sections": [
        {
            "title": "Approach",
            "section_html": "<p>We are <b>great</b> &amp; bold.</p>"
                            "<ul><li>One</li><li>Two</li></ul>"
                            "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>",
            "subsections": [
                {"title": "Details", "section_html": "<p>Deep dive.</p>",
                 "subsections": [{"title": "Security & Trust", "section_html": "<p>Zero trust.</p>"}]},
            ],
        },
    ],
    "has_past_performance": True,
    "past_performance": [{"award_title": "Big Win", "role": "Prime",
                          "description": "<p>Did <i>well</i>.</p>"}],
    "has_resumes": True,
    "resumes": [{"first_name": "Jane", "last_name": "Doe", "job_title": "Lead",
                 "professional_experience": [{"position": "Eng", "company_name": "Acme",
                                              "description": "<p>Built things.</p>"}],
                 "education": [{"degree": "BS", "institution_name": "MIT"}]}],
}


def _render(data):
    out = merge(build_example_template(), data)
    return docx.Document(io.BytesIO(out))


def test_full_generation():
    d = _render(SAMPLE)
    texts = [p.text for p in d.paragraphs]
    joined = "\n".join(texts)
    assert "Acme Corp" in joined
    assert "1  Approach" in joined          # auto-numbered section
    assert "1.1  Details" in joined         # auto-numbered subsection
    assert "1.1.1  Security & Trust" in joined  # ampersand preserved + depth 3
    assert "We are great & bold." in joined  # inline bold + decoded entity
    assert any(t.strip() == "\u2022\u00a0One" for t in texts)  # flattened bullet
    assert len(d.tables) == 1
    assert "Past Performance" in joined
    assert "Big Win" in joined
    assert "Jane Doe \u2014 Lead" in joined


def test_conditionals_hide_empty_blocks():
    d = _render({"metadata": {"title": "T"}, "sections": [{"title": "Only"}]})
    joined = "\n".join(p.text for p in d.paragraphs)
    assert "Past Performance" not in joined
    assert "Key Personnel" not in joined
    assert "1  Only" in joined


def test_empty_data_does_not_crash():
    d = _render({})
    assert len(d.paragraphs) > 0


if __name__ == "__main__":
    test_full_generation()
    test_conditionals_hide_empty_blocks()
    test_empty_data_does_not_crash()
    print("all tests passed")
